from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import json
import io
import re
import traceback
import os
import hmac
import hashlib
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import copy
from datetime import datetime

import sqlite3
from contextlib import contextmanager
from standardwebhooks.webhooks import Webhook, WebhookVerificationError

app = Flask(__name__)
CORS(app)

# ============================================================
# DODO PAYMENTS CONFIG
# ============================================================
DODO_WEBHOOK_SECRET = os.environ.get('DODO_WEBHOOK_SECRET', '')

PRODUCTS = {
    'monthly':  'pdt_0Nh18HtHXIP9Od1cy1DoE',
    'yearly':   'pdt_0Nh18Xr5AvKcgtPL3AYGT',
    'lifetime': 'pdt_0Nh18pGGx3eNDK1y4p6a0',
}
PRODUCT_TO_PLAN = {v: k for k, v in PRODUCTS.items()}

FREE_LIMIT = 3

# ============================================================
# PERSISTENT STORAGE — SQLite on Railway volume
# ============================================================
# Mount a Railway volume at /data so the DB survives redeploys.
# Falls back to local file if /data doesn't exist (e.g. local dev).
DB_DIR = '/data' if os.path.isdir('/data') else os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(DB_DIR, 'subscriptions.db')


def init_db():
    with get_db() as conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS paid_users (
                email TEXT PRIMARY KEY,
                active INTEGER NOT NULL DEFAULT 0,
                plan TEXT,
                since TEXT
            )
        ''')
        conn.commit()


@contextmanager
def get_db():
    conn = sqlite3.connect(DB_PATH)
    try:
        yield conn
    finally:
        conn.close()


def set_user_status(email, active, plan=None):
    email = email.lower()
    with get_db() as conn:
        conn.execute('''
            INSERT INTO paid_users (email, active, plan, since)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(email) DO UPDATE SET
                active = excluded.active,
                plan = COALESCE(excluded.plan, paid_users.plan),
                since = CASE WHEN excluded.active = 1 THEN excluded.since ELSE paid_users.since END
        ''', (email, 1 if active else 0, plan, datetime.now().isoformat()))
        conn.commit()


def get_user_status(email):
    email = email.lower()
    with get_db() as conn:
        row = conn.execute(
            'SELECT active, plan FROM paid_users WHERE email = ?', (email,)
        ).fetchone()
    if not row:
        return {'active': False, 'plan': None}
    return {'active': bool(row[0]), 'plan': row[1]}


init_db()

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_el(ns, tag):
    return etree.Element(f'{{{ns}}}{tag}')

def sub_el(parent, ns, tag):
    return etree.SubElement(parent, f'{{{ns}}}{tag}')

def make_run(text, italic=True, bold=False):
    r = make_el(MATH_NS, 'r')
    rpr = sub_el(r, MATH_NS, 'rPr')
    if not italic:
        sty = sub_el(rpr, MATH_NS, 'sty')
        sty.set(f'{{{MATH_NS}}}val', 'b' if bold else 'p')
    else:
        sty = sub_el(rpr, MATH_NS, 'sty')
        sty.set(f'{{{MATH_NS}}}val', 'bi' if bold else 'i')
    wrpr = sub_el(r, W_NS, 'rPr')
    rfonts = sub_el(wrpr, W_NS, 'rFonts')
    rfonts.set(f'{{{W_NS}}}ascii', 'Cambria Math')
    rfonts.set(f'{{{W_NS}}}hAnsi', 'Cambria Math')
    t = sub_el(r, MATH_NS, 't')
    t.text = text
    t.set(f'{{{W_NS}}}space', 'preserve')
    return r

def make_text_run(text):
    return make_run(text, italic=False, bold=False)

def make_frac(num_elements, den_elements):
    f = make_el(MATH_NS, 'f')
    fpr = sub_el(f, MATH_NS, 'fPr')
    ftype = sub_el(fpr, MATH_NS, 'type')
    ftype.set(f'{{{MATH_NS}}}val', 'bar')
    num = sub_el(f, MATH_NS, 'num')
    for el in num_elements:
        num.append(el)
    den = sub_el(f, MATH_NS, 'den')
    for el in den_elements:
        den.append(el)
    return f

def make_sup(base_elements, sup_elements):
    ssup = make_el(MATH_NS, 'sSup')
    e = sub_el(ssup, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    s = sub_el(ssup, MATH_NS, 'sup')
    for el in sup_elements:
        s.append(el)
    return ssup

def make_sub_el(base_elements, sub_elements):
    ssub = make_el(MATH_NS, 'sSub')
    e = sub_el(ssub, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    s = sub_el(ssub, MATH_NS, 'sub')
    for el in sub_elements:
        s.append(el)
    return ssub

def make_subsup(base_elements, sub_elements, sup_elements):
    ssubsup = make_el(MATH_NS, 'sSubSup')
    e = sub_el(ssubsup, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    sb = sub_el(ssubsup, MATH_NS, 'sub')
    for el in sub_elements:
        sb.append(el)
    sp = sub_el(ssubsup, MATH_NS, 'sup')
    for el in sup_elements:
        sp.append(el)
    return ssubsup

def make_sqrt(content_elements, degree_elements=None):
    rad = make_el(MATH_NS, 'rad')
    radpr = sub_el(rad, MATH_NS, 'radPr')
    if degree_elements is None:
        deghide = sub_el(radpr, MATH_NS, 'degHide')
        deghide.set(f'{{{MATH_NS}}}val', '1')
    deg = sub_el(rad, MATH_NS, 'deg')
    if degree_elements:
        for el in degree_elements:
            deg.append(el)
    e = sub_el(rad, MATH_NS, 'e')
    for el in content_elements:
        e.append(el)
    return rad

def make_accent(base_elements, accent_char='\u0302'):
    acc = make_el(MATH_NS, 'acc')
    accpr = sub_el(acc, MATH_NS, 'accPr')
    chr_el = sub_el(accpr, MATH_NS, 'chr')
    chr_el.set(f'{{{MATH_NS}}}val', accent_char)
    e = sub_el(acc, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    return acc

def make_delim(content_elements, beg='(', end=')'):
    d = make_el(MATH_NS, 'd')
    dpr = sub_el(d, MATH_NS, 'dPr')
    beg_el = sub_el(dpr, MATH_NS, 'begChr')
    beg_el.set(f'{{{MATH_NS}}}val', beg)
    end_el = sub_el(dpr, MATH_NS, 'endChr')
    end_el.set(f'{{{MATH_NS}}}val', end)
    e = sub_el(d, MATH_NS, 'e')
    for el in content_elements:
        e.append(el)
    return d

def make_func(func_name, arg_elements):
    func = make_el(MATH_NS, 'func')
    fname = sub_el(func, MATH_NS, 'fName')
    fname.append(make_run(func_name, italic=False))
    e = sub_el(func, MATH_NS, 'e')
    for el in arg_elements:
        e.append(el)
    return func

def make_nary(symbol, sub_els=None, sup_els=None, content_els=None):
    nary = make_el(MATH_NS, 'nary')
    narypr = sub_el(nary, MATH_NS, 'naryPr')
    chr_el = sub_el(narypr, MATH_NS, 'chr')
    chr_el.set(f'{{{MATH_NS}}}val', symbol)
    if sub_els is None:
        limLoc = sub_el(narypr, MATH_NS, 'limLoc')
        limLoc.set(f'{{{MATH_NS}}}val', 'undOvr')
    sub_e = sub_el(nary, MATH_NS, 'sub')
    if sub_els:
        for el in sub_els:
            sub_e.append(el)
    sup_e = sub_el(nary, MATH_NS, 'sup')
    if sup_els:
        for el in sup_els:
            sup_e.append(el)
    e = sub_el(nary, MATH_NS, 'e')
    if content_els:
        for el in content_els:
            e.append(el)
    return nary

LATEX_SYMBOLS = {
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\epsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ',
    r'\iota': 'ι', r'\kappa': 'κ', r'\lambda': 'λ', r'\mu': 'μ',
    r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π', r'\rho': 'ρ',
    r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ', r'\phi': 'φ',
    r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    r'\Alpha': 'Α', r'\Beta': 'Β', r'\Gamma': 'Γ', r'\Delta': 'Δ',
    r'\Epsilon': 'Ε', r'\Zeta': 'Ζ', r'\Eta': 'Η', r'\Theta': 'Θ',
    r'\Iota': 'Ι', r'\Kappa': 'Κ', r'\Lambda': 'Λ', r'\Mu': 'Μ',
    r'\Nu': 'Ν', r'\Xi': 'Ξ', r'\Pi': 'Π', r'\Rho': 'Ρ',
    r'\Sigma': 'Σ', r'\Tau': 'Τ', r'\Upsilon': 'Υ', r'\Phi': 'Φ',
    r'\Chi': 'Χ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
    r'\infty': '∞', r'\partial': '∂', r'\nabla': '∇',
    r'\pm': '±', r'\mp': '∓', r'\times': '×', r'\div': '÷',
    r'\cdot': '·', r'\circ': '∘', r'\bullet': '•',
    r'\leq': '≤', r'\geq': '≥', r'\ll': '≪', r'\gg': '≫',
    r'\neq': '≠', r'\ne': '≠', r'\approx': '≈', r'\equiv': '≡',
    r'\sim': '∼', r'\simeq': '≃', r'\propto': '∝',
    r'\in': '∈', r'\notin': '∉', r'\subset': '⊂', r'\supset': '⊃',
    r'\cup': '∪', r'\cap': '∩', r'\emptyset': '∅',
    r'\forall': '∀', r'\exists': '∃', r'\nexists': '∄',
    r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
    r'\Leftarrow': '⇐', r'\leftrightarrow': '↔', r'\Leftrightarrow': '⇔',
    r'\uparrow': '↑', r'\downarrow': '↓',
    r'\ldots': '…', r'\cdots': '⋯', r'\vdots': '⋮', r'\ddots': '⋱',
    r'\hbar': 'ℏ', r'\ell': 'ℓ', r'\Re': 'ℜ', r'\Im': 'ℑ',
    r'\aleph': 'ℵ', r'\wp': '℘',
    r'\oplus': '⊕', r'\otimes': '⊗', r'\odot': '⊙',
    r'\perp': '⊥', r'\parallel': '∥', r'\angle': '∠',
    r'\triangle': '△', r'\square': '□', r'\diamond': '◇',
    r'\star': '⋆', r'\dagger': '†', r'\ddagger': '‡',
    r'\langle': '⟨', r'\rangle': '⟩',
    r'\{': '{', r'\}': '}', r'\|': '‖',
    r'\%': '%', r'\$': '$', r'\#': '#', r'\&': '&',
    r'\quad': ' ', r'\qquad': '  ', r'\ ': ' ', r'\,': ' ',
    r'\mathbb{R}': 'ℝ', r'\mathbb{N}': 'ℕ', r'\mathbb{Z}': 'ℤ',
    r'\mathbb{Q}': 'ℚ', r'\mathbb{C}': 'ℂ',
}

def parse_latex(latex):
    latex = latex.strip()
    elements = []
    i = 0
    while i < len(latex):
        if latex[i] == '\\':
            matched = False
            for cmd in sorted(LATEX_SYMBOLS.keys(), key=len, reverse=True):
                if latex[i:].startswith(cmd):
                    after = i + len(cmd)
                    if after < len(latex) and latex[after].isalpha():
                        continue
                    elements.append(make_text_run(LATEX_SYMBOLS[cmd]))
                    i = after
                    if i < len(latex) and latex[i] == ' ':
                        i += 1
                    matched = True
                    break
            if not matched:
                m = re.match(r'\\([a-zA-Z]+)\*?', latex[i:])
                if m:
                    cmd_name = m.group(1)
                    after = i + m.end()
                    if cmd_name == 'frac':
                        num_group, after = _read_group(latex, after)
                        den_group, after = _read_group(latex, after)
                        num_els = parse_latex(num_group)
                        den_els = parse_latex(den_group)
                        elements.append(make_frac(num_els, den_els))
                        i = after
                    elif cmd_name == 'sqrt':
                        if after < len(latex) and latex[after] == '[':
                            end_bracket = latex.index(']', after)
                            degree_str = latex[after+1:end_bracket]
                            after = end_bracket + 1
                            content_group, after = _read_group(latex, after)
                            deg_els = parse_latex(degree_str)
                            content_els = parse_latex(content_group)
                            elements.append(make_sqrt(content_els, deg_els))
                        else:
                            content_group, after = _read_group(latex, after)
                            content_els = parse_latex(content_group)
                            elements.append(make_sqrt(content_els))
                        i = after
                    elif cmd_name in ('hat', 'vec', 'dot', 'ddot', 'tilde', 'bar', 'overline'):
                        accent_map = {
                            'hat': '\u0302', 'vec': '\u20d7', 'dot': '\u0307',
                            'ddot': '\u0308', 'tilde': '\u0303', 'bar': '\u0305',
                            'overline': '\u0305'
                        }
                        content_group, after = _read_group(latex, after)
                        content_els = parse_latex(content_group)
                        elements.append(make_accent(content_els, accent_map.get(cmd_name, '\u0302')))
                        i = after
                    elif cmd_name in ('left',):
                        if after < len(latex):
                            beg_char = latex[after]
                            if beg_char == '\\' and after+1 < len(latex):
                                beg_char = latex[after+1]
                                after += 2
                            else:
                                after += 1
                            right_pos = latex.find(r'\right', after)
                            if right_pos != -1:
                                inner = latex[after:right_pos]
                                after = right_pos + len(r'\right')
                                if after < len(latex):
                                    end_char = latex[after]
                                    if end_char == '\\' and after+1 < len(latex):
                                        end_char = latex[after+1]
                                        after += 2
                                    else:
                                        after += 1
                                else:
                                    end_char = ')'
                                inner_els = parse_latex(inner)
                                if beg_char == '(':
                                    elements.append(make_delim(inner_els, '(', ')'))
                                elif beg_char == '[':
                                    elements.append(make_delim(inner_els, '[', ']'))
                                elif beg_char == '{':
                                    elements.append(make_delim(inner_els, '{', '}'))
                                elif beg_char == '|':
                                    elements.append(make_delim(inner_els, '|', '|'))
                                else:
                                    elements.append(make_delim(inner_els, beg_char, end_char))
                                i = after
                            else:
                                elements.append(make_text_run('('))
                                i = after
                        else:
                            i += 1
                    elif cmd_name in ('sum', 'prod', 'int', 'oint', 'iint', 'iiint',
                                      'bigcup', 'bigcap', 'bigoplus', 'bigotimes'):
                        nary_map = {r'\sum': '∑', r'\prod': '∏', r'\int': '∫',
                                    r'\oint': '∮', r'\iint': '∬', r'\iiint': '∭',
                                    r'\bigcup': '⋃', r'\bigcap': '⋂',
                                    r'\bigoplus': '⊕', r'\bigotimes': '⊗'}
                        symbol = nary_map.get(f'\\{cmd_name}', '∑')
                        sub_els_n = None
                        sup_els_n = None
                        temp_i = after
                        if temp_i < len(latex) and latex[temp_i] == '_':
                            temp_i += 1
                            sg, temp_i = _read_group_or_char(latex, temp_i)
                            sub_els_n = parse_latex(sg)
                        if temp_i < len(latex) and latex[temp_i] == '^':
                            temp_i += 1
                            sg, temp_i = _read_group_or_char(latex, temp_i)
                            sup_els_n = parse_latex(sg)
                        elif sub_els_n and temp_i < len(latex) and latex[temp_i] == '_':
                            temp_i += 1
                            sg, temp_i = _read_group_or_char(latex, temp_i)
                            sup_els_n = parse_latex(sg)
                        content_g = ''
                        if temp_i < len(latex) and latex[temp_i] == '{':
                            content_g, temp_i = _read_group(latex, temp_i)
                        content_els_n = parse_latex(content_g) if content_g else []
                        elements.append(make_nary(symbol, sub_els_n, sup_els_n, content_els_n))
                        i = temp_i
                    elif cmd_name in ('sin', 'cos', 'tan', 'cot', 'sec', 'csc',
                                      'arcsin', 'arccos', 'arctan', 'sinh', 'cosh',
                                      'tanh', 'log', 'ln', 'exp', 'lim', 'max', 'min',
                                      'sup', 'inf', 'det', 'dim', 'ker', 'gcd'):
                        while after < len(latex) and latex[after] == ' ':
                            after += 1
                        if after < len(latex) and latex[after] == '{':
                            arg_group, after = _read_group(latex, after)
                            arg_els = parse_latex(arg_group)
                        elif after < len(latex):
                            ch = latex[after]
                            after += 1
                            arg_els = [make_text_run(ch)]
                        else:
                            arg_els = []
                        elements.append(make_func(cmd_name, arg_els))
                        i = after
                    elif cmd_name in ('mathbf', 'mathit', 'mathrm', 'mathsf',
                                      'mathtt', 'mathcal', 'mathbb', 'mathfrak',
                                      'text', 'textrm', 'textit', 'textbf',
                                      'operatorname'):
                        content_group, after = _read_group(latex, after)
                        elements.append(make_text_run(content_group))
                        i = after
                    elif cmd_name == 'begin':
                        env_group, after = _read_group(latex, after)
                        end_tag = f'\\end{{{env_group}}}'
                        end_pos = latex.find(end_tag, after)
                        if end_pos != -1:
                            inner = latex[after:end_pos]
                            after = end_pos + len(end_tag)
                            if env_group in ('matrix', 'pmatrix', 'bmatrix',
                                             'vmatrix', 'Vmatrix', 'cases', 'align',
                                             'align*', 'aligned'):
                                mat_el = parse_matrix_env(f'\\begin{{{env_group}}}{inner}\\end{{{env_group}}}')
                                if mat_el is not None:
                                    elements.append(mat_el)
                                else:
                                    elements.extend(parse_latex(inner))
                            else:
                                elements.extend(parse_latex(inner))
                        else:
                            elements.append(make_text_run(f'\\begin{{{env_group}}}'))
                        i = after
                    else:
                        elements.append(make_text_run(f'\\{cmd_name}'))
                        i = after
                else:
                    elements.append(make_text_run(latex[i]))
                    i += 1
        elif latex[i] == '^':
            i += 1
            sup_group, i = _read_group_or_char(latex, i)
            sup_els = parse_latex(sup_group)
            if elements:
                base = [elements.pop()]
                if i < len(latex) and latex[i] == '_':
                    i += 1
                    sub_group, i = _read_group_or_char(latex, i)
                    sub_els = parse_latex(sub_group)
                    elements.append(make_subsup(base, sub_els, sup_els))
                else:
                    elements.append(make_sup(base, sup_els))
            else:
                elements.append(make_sup([make_text_run('')], sup_els))
        elif latex[i] == '_':
            i += 1
            sub_group, i = _read_group_or_char(latex, i)
            sub_els = parse_latex(sub_group)
            if elements:
                base = [elements.pop()]
                if i < len(latex) and latex[i] == '^':
                    i += 1
                    sup_group, i = _read_group_or_char(latex, i)
                    sup_els = parse_latex(sup_group)
                    elements.append(make_subsup(base, sub_els, sup_els))
                else:
                    elements.append(make_sub_el(base, sub_els))
            else:
                elements.append(make_sub_el([make_text_run('')], sub_els))
        elif latex[i] == '{':
            group, i = _read_group(latex, i)
            group_els = parse_latex(group)
            elements.extend(group_els)
        elif latex[i] == '}':
            i += 1
        else:
            elements.append(make_text_run(latex[i]))
            i += 1
    return elements

def _find_matching_right(s, start):
    depth = 0
    i = start
    while i < len(s):
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            if depth == 0:
                return i
            depth -= 1
        i += 1
    return len(s)

def _read_group(s, pos):
    while pos < len(s) and s[pos] == ' ':
        pos += 1
    if pos >= len(s):
        return '', pos
    if s[pos] == '{':
        end = _find_matching_right(s, pos + 1)
        return s[pos+1:end], end + 1
    return s[pos], pos + 1

def _read_group_or_char(s, pos):
    while pos < len(s) and s[pos] == ' ':
        pos += 1
    if pos >= len(s):
        return '', pos
    if s[pos] == '{':
        return _read_group(s, pos)
    if s[pos] == '\\':
        m = re.match(r'\\([a-zA-Z]+)\*?', s[pos:])
        if m:
            end = pos + m.end()
            return s[pos:end], end
    return s[pos], pos + 1

def make_matrix(env, rows_data):
    m = make_el(MATH_NS, 'm')
    mpr = sub_el(m, MATH_NS, 'mPr')
    if env in ('pmatrix',):
        mbeg = sub_el(mpr, MATH_NS, 'begChr')
        mbeg.set(f'{{{MATH_NS}}}val', '(')
        mend = sub_el(mpr, MATH_NS, 'endChr')
        mend.set(f'{{{MATH_NS}}}val', ')')
    elif env in ('bmatrix',):
        mbeg = sub_el(mpr, MATH_NS, 'begChr')
        mbeg.set(f'{{{MATH_NS}}}val', '[')
        mend = sub_el(mpr, MATH_NS, 'endChr')
        mend.set(f'{{{MATH_NS}}}val', ']')
    elif env in ('vmatrix',):
        mbeg = sub_el(mpr, MATH_NS, 'begChr')
        mbeg.set(f'{{{MATH_NS}}}val', '|')
        mend = sub_el(mpr, MATH_NS, 'endChr')
        mend.set(f'{{{MATH_NS}}}val', '|')
    for row in rows_data:
        mr = sub_el(m, MATH_NS, 'mr')
        for cell in row:
            me = sub_el(mr, MATH_NS, 'e')
            cell_els = parse_latex(cell.strip())
            for el in cell_els:
                me.append(el)
    return m

def parse_matrix_env(latex):
    m = re.match(r'\\begin\{(\w+\*?)\}(.*?)\\end\{\1\}', latex, re.DOTALL)
    if not m:
        return None
    env = m.group(1)
    inner = m.group(2).strip()
    rows_raw = re.split(r'\\\\', inner)
    rows_data = []
    for row in rows_raw:
        row = row.strip()
        if not row:
            continue
        cells = re.split(r'(?<!\\)&', row)
        rows_data.append(cells)
    if not rows_data:
        return None
    return make_matrix(env, rows_data)

def build_omath(latex):
    omath = make_el(MATH_NS, 'oMath')
    elements = parse_latex(latex)
    for el in elements:
        omath.append(el)
    return omath

def insert_math(paragraph, latex):
    omath = build_omath(latex)
    paragraph._p.append(omath)

def add_block_formula(doc, latex):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath = build_omath(latex)
    p._p.append(omath)
    return p

def process_content(doc, content):
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            _code(doc, '\n'.join(code_lines))
            i += 1
            continue

        # Table
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _table_with_math(doc, table_lines)
            continue

        # Block formula $$...$$
        if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            latex = stripped[2:-2].strip()
            add_block_formula(doc, latex)
            i += 1
            continue

        if stripped == '$$':
            latex_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                latex_lines.append(lines[i])
                i += 1
            add_block_formula(doc, '\n'.join(latex_lines))
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^[\*\-\+]\s+(.*)', stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            _fmt(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style='List Number')
            _fmt(p, text)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[\-\*\_]{3,}$', stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blank line
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _text_math(doc, stripped) if ('$$' in stripped or (stripped.count('$') >= 2)) else _fmt(p, stripped)
        if '$' in stripped:
            p.clear()
            _text_math_inline(p, stripped)
        i += 1

def _text_math(doc, text):
    parts = re.split(r'\$\$(.*?)\$\$', text, flags=re.DOTALL)
    if len(parts) == 1:
        p = doc.add_paragraph()
        _fmt(p, text)
        return
    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            if part.strip():
                p = doc.add_paragraph()
                _fmt(p, part.strip())
        else:
            add_block_formula(doc, part.strip())

def _text_math_inline(para, text):
    parts = re.split(r'\$(.*?)\$', text)
    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            if part:
                _fmt(para, part)
        else:
            insert_math(para, part.strip())

def _add_cell_content_with_math(cell, text):
    p = cell.paragraphs[0]
    if '$' in text:
        _text_math_inline(p, text)
    else:
        _fmt(p, text)

def _table_with_math(doc, tlines):
    rows = []
    for line in tlines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\|[\s\-\|:]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = ''
                _add_cell_content_with_math(cell, cell_text)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

def _has_emoji(text):
    import unicodedata
    for ch in text:
        try:
            if unicodedata.category(ch) in ('So', 'Cs'):
                return True
            name = unicodedata.name(ch, '')
            if 'EMOJI' in name or 'PICTOGRAPH' in name:
                return True
        except Exception:
            pass
    return False

def _add_run_with_emoji(para, text, bold=False, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if _has_emoji(text):
        rpr = run._r.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), 'Segoe UI Emoji')
        rfonts.set(qn('w:hAnsi'), 'Segoe UI Emoji')
        rpr.append(rfonts)

def _fmt(para, text):
    text = re.sub(r'\*\*\*(.*?)\*\*\*', lambda m: f'\x00BOLDITALIC\x01{m.group(1)}\x02', text)
    text = re.sub(r'\*\*(.*?)\*\*', lambda m: f'\x00BOLD\x01{m.group(1)}\x02', text)
    text = re.sub(r'\*(.*?)\*', lambda m: f'\x00ITALIC\x01{m.group(1)}\x02', text)
    text = re.sub(r'`([^`]+)`', lambda m: f'\x00CODE\x01{m.group(1)}\x02', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    parts = re.split(r'(\x00(?:BOLD|ITALIC|BOLDITALIC|CODE)\x01.*?\x02)', text)
    for part in parts:
        if part.startswith('\x00BOLDITALIC\x01'):
            content = part[len('\x00BOLDITALIC\x01'):-1]
            _add_run_with_emoji(para, content, bold=True, italic=True)
        elif part.startswith('\x00BOLD\x01'):
            content = part[len('\x00BOLD\x01'):-1]
            _add_run_with_emoji(para, content, bold=True)
        elif part.startswith('\x00ITALIC\x01'):
            content = part[len('\x00ITALIC\x01'):-1]
            _add_run_with_emoji(para, content, italic=True)
        elif part.startswith('\x00CODE\x01'):
            content = part[len('\x00CODE\x01'):-1]
            run = para.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif part:
            _add_run_with_emoji(para, part)

def _code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

def _img(doc, src, alt=''):
    pass

# ============================================================
# HEALTH CHECK
# ============================================================
@app.route('/', methods=['GET'])
@app.route('/api/export-chat', methods=['GET'])
def health():
    return jsonify({'status': 'OK', 'version': '5.0-flask'})

# ============================================================
# EXPORT CHAT (existing)
# ============================================================
@app.route('/api/export-chat', methods=['POST'])
def export_chat():
    try:
        data = request.get_json(force=True)
        messages = data.get('messages', [])
        title = data.get('title', 'Gemini Chat')

        if not messages:
            return jsonify({'error': 'No messages'}), 400

        doc = Document()

        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp = doc.add_paragraph()
        dr = dp.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))
        dr.font.size = Pt(10)
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for i, msg in enumerate(messages):
            role = msg.get('role', 'user')
            content = msg.get('content', '')

            rp = doc.add_paragraph()
            rr = rp.add_run('You' if role == 'user' else 'Gemini')
            rr.bold = True
            rr.font.size = Pt(14)
            rr.font.color.rgb = RGBColor(33, 150, 243) if role == 'user' else RGBColor(76, 175, 80)

            process_content(doc, content)

            if i < len(messages) - 1:
                sp = doc.add_paragraph()
                sr = sp.add_run('─' * 60)
                sr.font.color.rgb = RGBColor(200, 200, 200)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='gemini-chat.docx'
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500


# ============================================================
# DODO PAYMENTS — WEBHOOK & STATUS (below)
# ============================================================


# ============================================================
# DODO PAYMENTS — WEBHOOK
# ============================================================
@app.route('/api/webhook', methods=['POST'])
def webhook():
    try:
        payload = request.get_data()
        webhook_headers = {
            'webhook-id': request.headers.get('webhook-id', ''),
            'webhook-signature': request.headers.get('webhook-signature', ''),
            'webhook-timestamp': request.headers.get('webhook-timestamp', ''),
        }

        # Signature verification — required, do not disable
        if not DODO_WEBHOOK_SECRET:
            print('❌ DODO_WEBHOOK_SECRET not set — rejecting webhook')
            return jsonify({'error': 'Webhook secret not configured'}), 500

        try:
            wh = Webhook(DODO_WEBHOOK_SECRET)
            wh.verify(payload, webhook_headers)
        except WebhookVerificationError as e:
            print(f'❌ Invalid webhook signature: {e}')
            return jsonify({'error': 'Invalid signature'}), 400

        event = request.get_json(force=True)
        event_type = event.get('type', '')
        data = event.get('data', {})

        product_id = data.get('product_id') or (
            (data.get('product_cart') or [{}])[0].get('product_id')
        )
        plan = PRODUCT_TO_PLAN.get(product_id, product_id)

        if event_type in ('payment.succeeded', 'subscription.active'):
            customer = data.get('customer', {})
            email = customer.get('email', '').lower()
            if email:
                set_user_status(email, active=True, plan=plan)
                print(f'✅ Pro activated: {email} ({plan})')

        if event_type in ('subscription.cancelled', 'subscription.expired', 'subscription.failed'):
            customer = data.get('customer', {})
            email = customer.get('email', '').lower()
            if email:
                set_user_status(email, active=False, plan=plan)
                print(f'❌ Pro deactivated: {email}')

        return jsonify({'ok': True})

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ============================================================
# CHECK USER STATUS
# ============================================================
@app.route('/api/status', methods=['GET'])
def check_status():
    email = request.args.get('email', '').lower()
    if not email:
        return jsonify({'active': False})

    return jsonify(get_user_status(email))


# ============================================================
# SUCCESS PAGE (after payment)
# ============================================================
@app.route('/success', methods=['GET'])
def success_page():
    return '''
    <html>
    <head>
        <title>Payment Successful!</title>
        <style>
            body { font-family: -apple-system, sans-serif; display: flex;
                   align-items: center; justify-content: center;
                   height: 100vh; margin: 0; background: #f8f9fa; }
            .card { background: white; padding: 48px; border-radius: 16px;
                    text-align: center; box-shadow: 0 4px 24px rgba(0,0,0,0.1); }
            h1 { color: #34a853; margin-bottom: 8px; }
            p { color: #5f6368; margin-bottom: 24px; }
            .badge { background: #e8f5e9; color: #2e7d32;
                     padding: 8px 20px; border-radius: 20px;
                     font-weight: 600; font-size: 14px; }
        </style>
    </head>
    <body>
        <div class="card">
            <div style="font-size:64px">🎉</div>
            <h1>You're now Pro!</h1>
            <p>Your account has been activated. Close this tab and enjoy unlimited exports!</p>
            <div class="badge">✓ Pro activated</div>
        </div>
    </body>
    </html>
    '''


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port)
